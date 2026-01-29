import json
from pathlib import Path
from typing import List, Dict, Any
import pandas as pd
import csv

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
try:
    import docx
except ImportError:
    docx = None

from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from src.vectorstore import VectorStoreClient
from src.security import validate_sensitivity

CHUNK_SIZE = 1000  
CHUNK_OVERLAP = 200 

def extract_text_from_pdf(path: str) -> List[Dict[str, Any]]:
    if fitz is None:
        raise RuntimeError("PyMuPDF not installed")
    doc = fitz.open(path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        pages.append({"page": i+1, "text": text, "source_type": "pdf"})
    return pages

def extract_text_from_docx(path: str) -> List[Dict[str, Any]]:
    if docx is None:
        raise RuntimeError("python-docx not installed")
    d = docx.Document(path)
    full_text = []
    for p in d.paragraphs:
        if p.text.strip():
            full_text.append(p.text)
    return [{"page": 1, "text": "\n".join(full_text), "source_type": "docx"}]

def preprocess_table(path: str) -> pd.DataFrame:
    """
    Preparing CSV/XLS/XLSX files for ingestion:
    - Determines the correct delimiter for CSV files
    - Cleans column names
    - Fills NaN values ​​with empty strings
    """
    ext = path.lower()
    try:
        if ext.endswith(".csv"):
            # determine the delimiter using csv.Sniffer.
            with open(path, "r", encoding="utf-8-sig") as f:
                sample = f.read(2048)
                sniffer = csv.Sniffer()
                dialect = sniffer.sniff(sample)
                sep = dialect.delimiter
            df = pd.read_csv(path, sep=sep, encoding="utf-8-sig", low_memory=False)

        else:
            engine = "openpyxl" if ext.endswith(".xlsx") else "xlrd"
            df = pd.read_excel(path, engine=engine)
    except Exception as e:
        raise ValueError(f"Не удалось прочитать таблицу {path}: {e}")

    # Cleaning
    df.columns = [str(c).strip().replace(" ", "_") for c in df.columns]

    # filling the NaN values ​​with empty strings
    df = df.fillna("")

    # convert ALL columns to strings so that Parquet doesn't crash on mixed data types.
    for col in df.columns:
        df[col] = df[col].astype(str)

    return df

def process_structured_table(path: str, doc_id: str) -> List[Dict[str, Any]]:
    """
    Saves the table in parquet format and creates a text description of the schema.
    """
    print(f"[INGEST] Processing structured table: {path} with doc_id={doc_id}")

    df = preprocess_table(path)

    tables_dir = Path("data/tables")
    tables_dir.mkdir(parents=True, exist_ok=True)
    table_path = (tables_dir / f"{doc_id}.parquet").resolve()

    try:
        df.to_parquet(table_path, index=False)
    except Exception as e:
        raise RuntimeError(
            f"❌ Failed to save parquet file.\n"
            f"Install dependency: pip install pyarrow\n"
            f"Original error: {e}"
        )

    print(f"[INGEST] ✅ Table saved to {table_path}")

    summary = (
        f"TABLE DOCUMENT ID: {doc_id}\n"
        f"TYPE: Structured Data\n"
        f"ROWS: {len(df)}\n"
        f"COLUMNS: {list(df.columns)}\n"
        f"SAMPLE DATA:\n{df.head(3).to_markdown(index=False)}"
    )

    return [{
        "page": 1,
        "text": summary,
        "source_type": "structured_table",
        "file_path": str(table_path)
    }]

def create_documents(pages: List[Dict[str, Any]], doc_id: str, file_sensitivity: str) -> List[Document]:
    """
    Splits text and creates LlamaIndex Document objects (Nodes).
    """
    # LlamaIndex splitter
    splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    
    docs_to_insert = []
    global_counter = 0

    for p in pages:
        raw_text = (p.get("text") or "").strip()
        if not raw_text:
            continue

        # If it's our custom table markdown, we treat the whole batch as one chunk
        if p.get("source_type") == "structured_table":
            text_chunks = [raw_text]

        else:
            text_chunks = splitter.split_text(raw_text)

        page_no = p.get("page", 1)

        for part in text_chunks:
            chunk_id = f"{doc_id}_p{page_no}_{global_counter}"
            
            metadata = {
                "doc_id": doc_id,
                "page": page_no,
                "chunk_id": chunk_id,
                "source_type": p.get("source_type", "unknown"),
                "sensitivity": file_sensitivity
            }
            # pass parquet path to RAG
            if p.get("file_path"):
                metadata["table_path"] = p["file_path"] # Standardize key name
                metadata["file_path"] = p["file_path"]  # Keep for backward compatibility
            
            if "row_start" in p:
                metadata["row_idx"] = f"{p['row_start']}-{p['row_end']}"

            # Create LlamaIndex Document (which will become a Node in the index)
            # We explicitly set doc_id so we can overwrite it later (upsert logic)
                doc = Document(text=part, metadata=metadata)
                doc.doc_id = chunk_id
                doc.metadata["ref_doc_id"] = doc_id

                docs_to_insert.append(doc)
                global_counter += 1
    
    return docs_to_insert

def ingest_file(path: str, doc_id: str, sensitivity: str):
    path_obj = Path(path)
    if not path_obj.exists():
         raise FileNotFoundError(f"{path} does not exist.")
         
    handlers = {
        ".pdf":  lambda p: extract_text_from_pdf(str(p)),
        ".docx": lambda p: extract_text_from_docx(str(p)),
        ".txt":  lambda p: [{"page": 1, "text": p.read_text(encoding="utf-8"), "source_type": "text"}],
        ".md":   lambda p: [{"page": 1, "text": p.read_text(encoding="utf-8"), "source_type": "text"}],
        ".csv":  lambda p: process_structured_table(str(p), doc_id),
        ".xlsx": lambda p: process_structured_table(str(p), doc_id),
        ".xls":  lambda p: process_structured_table(str(p), doc_id),
    }

    ext = path_obj.suffix.lower()
    handler = handlers.get(ext)
    if not handler:
        raise RuntimeError(f"Unsupported extension: {ext}")
    
    pages = handler(path_obj)
    clean_sensitivity = validate_sensitivity(sensitivity)
    
    # Convert to LlamaIndex Documents
    nodes = create_documents(pages, doc_id, clean_sensitivity)
    
    # Upsert to VectorStore
    vs_client = VectorStoreClient()
    vs_client.upsert_document(doc_id, nodes)
    
    # Manifest saving (Metadata)
    manifests_dir = Path("data") / "manifests"
    manifests_dir.mkdir(parents=True, exist_ok=True)
    table_path = str(Path("data/tables") / f"{doc_id}.parquet")
    man_path = Path("data") / f"{doc_id}_manifest.json"
    man_path.parent.mkdir(exist_ok=True)
    man_path.write_text(json.dumps({
        "doc_id": doc_id,
        "sensitivity": clean_sensitivity,
        "num_chunks": len(nodes),
        "table_path": str(table_path) #to find the necessary tables
    }), encoding="utf-8")
